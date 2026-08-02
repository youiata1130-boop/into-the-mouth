from pathlib import Path
from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "rules"
ASSET = ROOT / "docs" / "rulebook" / "assets"
HERO = ASSET / "whale-mouth-rulebook.png"
FISH2 = ROOT / "public" / "assets" / "creatures" / "fish-2-sardine.png"
FISH3 = ROOT / "public" / "assets" / "creatures" / "fish-3.png"
FISH4 = ROOT / "public" / "assets" / "creatures" / "creature-4-octopus.png"
FISH6 = ROOT / "public" / "assets" / "creatures" / "fish-5-shark.png"
POISON = ROOT / "public" / "assets" / "creatures" / "fish-poison.png"

NAVY = "12354A"
BLUE = "1C7890"
TEAL = "45A7A1"
CREAM = "FFF8E8"
CORAL = "E96D55"
YELLOW = "F1BB4D"
INK = "18333F"
MUTED = "55727C"
PALE = "E8F5F3"
WHITE = "FFFFFF"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tcMar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def no_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "nil")
        borders.append(el)


def set_cell_width(cell, mm):
    tcPr = cell._tc.get_or_add_tcPr()
    tcW = tcPr.find(qn("w:tcW"))
    if tcW is None:
        tcW = OxmlElement("w:tcW")
        tcPr.append(tcW)
    tcW.set(qn("w:w"), str(int(mm / 25.4 * 1440)))
    tcW.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_font(run, size=9.2, bold=False, color=INK, name="Yu Gothic"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def para(container, text="", size=9.2, bold=False, color=INK, align=None, before=0, after=3, line=1.12):
    p = container.add_paragraph() if hasattr(container, "add_paragraph") else container.paragraphs[0]
    if text:
        set_font(p.add_run(text), size, bold, color)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    return p


def clear_cell(cell):
    p = cell.paragraphs[0]
    p.clear()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    return p


def bullet(container, text, color=INK, after=1.6):
    p = container.add_paragraph(style="Rule Bullet")
    p.paragraph_format.space_after = Pt(after)
    set_font(p.add_run(text), 8.7, False, color)
    return p


def heading(container, number, title, color=BLUE, before=5, after=3):
    p = container.add_paragraph(style="Rule Heading")
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    r = p.add_run(f"{number}  {title}")
    set_font(r, 13.2, True, color)
    return p


def callout(doc, label, text, fill=CREAM, accent=CORAL):
    t = doc.add_table(rows=1, cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_cell_width(t.cell(0, 0), 25)
    set_cell_width(t.cell(0, 1), 154)
    shade(t.cell(0, 0), accent)
    shade(t.cell(0, 1), fill)
    for c in t.rows[0].cells:
        margins(c, 100, 140, 100, 140)
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = clear_cell(t.cell(0, 0)); set_font(p.add_run(label), 9, True, WHITE); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = clear_cell(t.cell(0, 1)); set_font(p.add_run(text), 8.7, True, INK)
    no_borders(t)
    para(doc, "", after=1)


def icon_row(doc):
    t = doc.add_table(rows=1, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    labels = [(FISH2, "魚2 ×3"), (FISH3, "魚3 ×3"), (FISH4, "魚4 ×1"), (FISH6, "魚6 ×1"), (POISON, "毒魚 ×1")]
    for i, (path, label) in enumerate(labels):
        c = t.cell(0, i); set_cell_width(c, 35.8); margins(c, 40, 60, 50, 60); shade(c, PALE if i < 4 else CREAM)
        p = clear_cell(c); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Mm(14), height=Mm(10))
        p = c.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after = Pt(0)
        set_font(p.add_run(label), 7.8, True, CORAL if i == 4 else NAVY)
    no_borders(t)


def step_strip(doc):
    t = doc.add_table(rows=1, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    steps = [("1", "親が口を開く"), ("2", "子がカードを出す"), ("3", "捕食・特殊効果"), ("4", "閉じる／逃げる")]
    for i, (n, txt) in enumerate(steps):
        c=t.cell(0,i); set_cell_width(c,44.7); margins(c,90,90,90,90); shade(c, [NAVY,BLUE,TEAL,CORAL][i])
        p=clear_cell(c); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(n+"\n"+txt), 8.3, True, WHITE)
    no_borders(t)


def add_page_field(paragraph):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    separate = OxmlElement("w:fldChar"); separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t"); text.text = "1"
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])
    set_font(run, 7.5, False, MUTED)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run("口に入る｜ルールガイド  "), 7.5, False, MUTED)
    add_page_field(p)
    set_font(p.add_run("/2"), 7.5, False, MUTED)


def build():
    OUT.mkdir(parents=True, exist_ok=True)
    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Mm(210); sec.page_height=Mm(297)
    sec.top_margin=Mm(9); sec.bottom_margin=Mm(10); sec.left_margin=Mm(14); sec.right_margin=Mm(14)
    sec.header_distance=Mm(4); sec.footer_distance=Mm(4)
    add_footer(sec)

    styles=doc.styles
    normal=styles["Normal"]
    normal.font.name="Yu Gothic"; normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size=Pt(9.2); normal.font.color.rgb=RGBColor.from_string(INK)
    normal.paragraph_format.space_after=Pt(3); normal.paragraph_format.line_spacing=1.12
    bs=styles.add_style("Rule Bullet", 1)
    bs.font.name="Yu Gothic"; bs._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic"); bs.font.size=Pt(8.7)
    bs.paragraph_format.left_indent=Mm(4.5); bs.paragraph_format.first_line_indent=Mm(-3.4); bs.paragraph_format.space_after=Pt(1.6)
    # real bullet numbering
    numPr=OxmlElement("w:numPr"); ilvl=OxmlElement("w:ilvl"); ilvl.set(qn("w:val"),"0"); numId=OxmlElement("w:numId"); numId.set(qn("w:val"),"1"); numPr.extend([ilvl,numId]); bs._element.get_or_add_pPr().append(numPr)
    hs=styles.add_style("Rule Heading", 1); hs.font.name="Yu Gothic"; hs._element.rPr.rFonts.set(qn("w:eastAsia"),"Yu Gothic")

    # PAGE 1
    title=doc.add_paragraph(); title.alignment=WD_ALIGN_PARAGRAPH.CENTER; title.paragraph_format.space_after=Pt(0)
    set_font(title.add_run("口に入る"), 27, True, NAVY)
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER; sub.paragraph_format.space_after=Pt(10)
    set_font(sub.add_run("魚を食べ、逃げ、毒で出し抜く  3〜6人用カードゲーム"), 9.4, True, NAVY)
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(4)
    p.add_run().add_picture(str(HERO), width=Mm(182), height=Mm(48))

    callout(doc,"勝ち方","全員が1回ずつ親を担当。最後に合計得点が最も高い人が勝ち。同点は全員勝利です。",PALE,TEAL)
    heading(doc,"1","ゲームの準備",before=4)
    bullet(doc,"親を1人決め、ほかの人は子になります。親ラウンドは必ず3トライです。")
    bullet(doc,"各子は自分の9枚をよく混ぜて山札にし、上から3枚を表向きに並べます。魚6には魚5のサメの絵を使います。")
    icon_row(doc)
    bullet(doc,"公開カードを使ったら山札から即補充。山札・公開カード・使用済みカードは3トライを通して引き継ぎます。")
    bullet(doc,"逃げる専用カードは山札に混ぜず隣へ置き、各トライ1回だけ使います。通常カードがなくても未使用なら使えます。")

    heading(doc,"2","1トライの流れ")
    step_strip(doc)
    para(doc,"",after=1)
    bullet(doc,"箱の最初に、親の餌「1」を入れます。これは親自身のカードなので親の得点には入りません。")
    bullet(doc,"親が口を開けたら、子は好きな公開カードを表向きで出せます。順番は固定ではありません。")
    bullet(doc,"魚が入るたびに捕食を解決します。子は条件を満たせば、山札の隣の専用カードを使って逃げられます。")
    bullet(doc,"親が口を閉じる、または逃げが成功するとトライ終了。入った順に公開して得点します。")
    callout(doc,"箱の見方","カードは入った順に処理。現物では下のカードほど先に入ったカードです。口が開いている間は一番上だけ見えます。",CREAM,YELLOW)
    heading(doc,"3","親ラウンドとゲーム終了")
    bullet(doc,"1人の親につき3トライ。3トライ後、次の人が親になり、子の山札を9枚に戻して混ぜ直します。")
    bullet(doc,"全員が親を終えたらゲーム終了。得点を合計して勝者を決めます。")

    doc.add_page_break()

    # PAGE 2
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(1)
    set_font(p.add_run("口の中で起こること"), 21, True, NAVY)
    para(doc,"捕食・魚群・逃げる・毒魚を押さえれば、すぐに遊べます。",9,False,MUTED,after=7)

    heading(doc,"4","捕食：強い魚が食べる。同点は後出し勝ち",before=0)
    bullet(doc,"新しい魚を出したら、直前のカードから逆向きに強さを比べます。強い魚が弱い魚を食べます。")
    bullet(doc,"先に強い魚がいる場所へ弱い魚を出すと、新しい魚が食べられます。同じ強さなら、後から出した魚が先にいた魚を食べます。")
    bullet(doc,"食べた魚がさらにカードを食べていた場合、その獲物もまとめて大きい魚へ引き継ぎます。")
    callout(doc,"例","餌1 → 魚2 → 魚4 の順なら、魚4が魚2と餌1をまとめて食べます。",PALE,BLUE)

    heading(doc,"5","魚群：同じ2・3を重ねて強くする")
    bullet(doc,"公開中の同じ魚2、または同じ魚3を2枚重ねると2匹の群れ。さらに同種を1枚追加して3匹にできます。")
    t=doc.add_table(rows=3, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    data=[("魚群","2匹","3匹"),("魚2","強さ4","強さ6"),("魚3","強さ6","強さ9")]
    for r,row in enumerate(data):
        for c,val in enumerate(row):
            cell=t.cell(r,c); set_cell_width(cell,[42,68,68][c]); margins(cell,70,100,70,100); shade(cell,NAVY if r==0 else (PALE if r==1 else CREAM)); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=clear_cell(cell); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; set_font(p.add_run(val),8.5,True,WHITE if r==0 else INK)
    no_borders(t); set_repeat_table_header(t.rows[0])
    bullet(doc,"3匹の群れは1枚の魚として扱い、分けたり重ね直したりできません。すべての判定は合計の強さで行います。",after=2)

    heading(doc,"6","逃げる：食べた分を確定得点に")
    bullet(doc,"自分の魚がほかの数字カードを食べていれば、その子には「逃げる権利」があります。")
    bullet(doc,"山札の隣に置いた専用カードを使います。成功すると、その魚が食べた他人の数字カードの合計を得点し、トライ終了です。逃げる魚自身は得点にしません。")
    bullet(doc,"専用カードは1トライ1回。候補が複数なら最後に出した有効な魚が対象です。権利なしで使うと不発ですが、そのトライでは再使用できません。次のトライで戻ります。")

    heading(doc,"7","毒魚：得点の権利を奪う")
    bullet(doc,"有効な毒魚は、別の子が次に出した魚1枚を毒魚の持ち主の確定得点にします。毒魚を出した本人が続けて魚を出すと、その魚は無効です。")
    bullet(doc,"別の毒魚が出ると権利は後の毒魚へ移ります。毒魚は魚1枚に発動すると無効になります。")
    bullet(doc,"親は口が開いていて毒魚が有効な間、時間制限なく取り除けます。残したまま口を閉じると、毒魚の持ち主が10点、親は0点でトライ終了です。")
    callout(doc,"重要","有効な毒魚、成功した逃げるカード、毒魚で得点化済みの魚に当たると捕食は止まります。",CREAM,CORAL)

    heading(doc,"8","得点早見",before=4)
    t=doc.add_table(rows=4, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER; t.autofit=False
    rows=[("だれ","いつ","得点"),("子（逃げる）","逃げ成功","対象魚が食べた他人の数字合計"),("子（毒魚）","次の魚／親が閉じる","対象魚の確定得点／10点"),("親","口を閉じる","残った他人の有効な数字合計")]
    widths=[35,45,98]
    for r,row in enumerate(rows):
        for c,val in enumerate(row):
            cell=t.cell(r,c); set_cell_width(cell,widths[c]); margins(cell,70,90,70,90); shade(cell,NAVY if r==0 else (PALE if r%2 else WHITE)); cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p=clear_cell(cell); set_font(p.add_run(val),7.8, r==0 or c==0, WHITE if r==0 else INK)
    set_repeat_table_header(t.rows[0]); no_borders(t)
    para(doc,"親得点には親自身の餌、毒魚、逃げるカード、毒魚で得点化済み／無効化された魚を含みません。口を開けてすぐ閉じれば0点です。",7.7,True,MUTED,after=0)

    path=OUT / "into-the-mouth-rules-ja.docx"
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build())
